"""Word (.docx) export via python-docx.

Depends on the ``docx`` optional extra (python-docx).
Designed for teachers who want to edit seating plans further in Word.
"""

from __future__ import annotations

from pathlib import Path

from seattrellis.exporters.presentation import (
    student_detail_fields,
    student_display_names,
)
from seattrellis.exporters.print_html import (
    _default_privacy,
    _recommendation_text,
    _text,
    _validate_template,
)
from seattrellis.models.candidate import CandidatePlan
from seattrellis.models.snapshot import SeatingSnapshot
from seattrellis.service_types import PageOptions, PrivacyOptions, normalize_export_locale


def export_docx(
    snapshot: SeatingSnapshot,
    output: str | Path,
    *,
    template: str = "public",
    privacy: PrivacyOptions | None = None,
    candidate: CandidatePlan | None = None,
    page: PageOptions | None = None,
    locale: str = "zh",
) -> Path:
    """Export a seating snapshot as a .docx file.

    Parameters
    ----------
    snapshot:
        The seating snapshot.
    output:
        Output ``.docx`` path.
    template:
        One of ``"public"``, ``"teacher"``, ``"report"``.
    privacy:
        Privacy options.
    candidate:
        Candidate plan for the ``"report"`` template.
    """
    template = _validate_template(template)
    locale = normalize_export_locale(locale)
    if template == "report" and candidate is None:
        raise ValueError("The report template requires a candidate plan.")

    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Inches, Pt  # type: ignore[import-untyped]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover
        from seattrellis.optional import MissingOptionalDependencyError

        raise MissingOptionalDependencyError("Word export", "docx") from exc

    if privacy is None:
        privacy = _default_privacy(template)

    page = page or PageOptions()
    doc = Document()
    section = doc.sections[0]
    if page.orientation == "landscape":
        from docx.enum.section import WD_ORIENT  # type: ignore[import-untyped]

        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
    section.top_margin = Inches(page.margin_mm / 25.4)
    section.bottom_margin = Inches(page.margin_mm / 25.4)
    section.left_margin = Inches(page.margin_mm / 25.4)
    section.right_margin = Inches(page.margin_mm / 25.4)

    # Title
    title = doc.add_heading(snapshot.layout.name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24 * page.scale)

    # Meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"{_text('generated_at', locale)}: {snapshot.created_at}"
    )
    meta_run.font.size = Pt(9 * page.scale)
    meta_run.font.color.rgb = None  # default

    # Candidate info if available
    if candidate is not None:
        cand_para = doc.add_paragraph()
        cand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cand_run = cand_para.add_run(
            f"{_text('candidate', locale)}: {candidate.candidate_id}  |  "
            f"{_text('total_score', locale)}: {candidate.total_score:.1f}"
        )
        cand_run.font.size = Pt(10 * page.scale)
        cand_run.bold = True

    doc.add_paragraph()  # spacer

    # Seat table
    min_row, max_row, min_col, max_col = _bounds(snapshot)
    seat_by_pos = {(s.row, s.col): s for s in snapshot.layout.seats}
    assign_by_seat = {a.seat_id: a for a in snapshot.assignments}

    table = doc.add_table(rows=max_row - min_row + 1, cols=max_col - min_col + 1)
    table.style = "Table Grid"

    display_names = student_display_names(snapshot, privacy, locale)

    for r in range(min_row, max_row + 1):
        for c in range(min_col, max_col + 1):
            cell = table.cell(r - min_row, c - min_col)
            seat = seat_by_pos.get((r, c))
            if seat is None:
                cell.text = ""
                continue
            a = assign_by_seat.get(seat.seat_id)
            name = display_names.get(a.student_key, "") if a else ""
            cell.text = name if (name and seat.enabled) else seat.seat_id
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10 * page.scale)

    # Teacher section
    if template == "teacher":
        student_by_key = {s.key: s for s in snapshot.students}
        detail_headers = [
            header
            for header, _value in student_detail_fields(None, privacy, locale)
        ]
        doc.add_heading(_text("student_details", locale), level=2)
        detail_table = doc.add_table(
            rows=len(snapshot.assignments) + 1,
            cols=2 + len(detail_headers),
        )
        detail_table.style = "Table Grid"
        headers = [
            _text("seat", locale),
            _text("name", locale),
            *detail_headers,
        ]
        for i, h in enumerate(headers):
            detail_table.cell(0, i).text = h
            _set_cell_font_size(detail_table.cell(0, i), Pt(10 * page.scale))
        for i, a in enumerate(snapshot.assignments):
            stu = student_by_key.get(a.student_key)
            detail_table.cell(i + 1, 0).text = a.seat_id
            detail_table.cell(i + 1, 1).text = display_names[a.student_key]
            _set_cell_font_size(
                detail_table.cell(i + 1, 0), Pt(10 * page.scale)
            )
            _set_cell_font_size(
                detail_table.cell(i + 1, 1), Pt(10 * page.scale)
            )
            for column, (_header, value) in enumerate(
                student_detail_fields(stu, privacy, locale),
                start=2,
            ):
                detail_table.cell(i + 1, column).text = value
                _set_cell_font_size(
                    detail_table.cell(i + 1, column), Pt(10 * page.scale)
                )

    # Report section
    if template == "report" and candidate is not None:
        doc.add_heading(_text("report_title", locale), level=2)
        b = candidate.score.breakdown
        for dim_name, dim_score in [
            (_text("fair_rotation", locale), b.fair_rotation_score),
            (_text("neighbor_avoidance", locale), b.avoid_recent_neighbors_score),
            (_text("score_mixing", locale), b.score_balance_score),
            (_text("height_preference", locale), b.height_preference_score),
            (_text("vision_preference", locale), b.vision_preference_score),
            (_text("diversity", locale), b.diversity_score),
            (_text("stability", locale), b.stability_score),
        ]:
            score_str = f"{dim_score.score:.1f}" if dim_score.score is not None else "n/a"
            paragraph = doc.add_paragraph(
                f"{dim_name}: {score_str} (weight={dim_score.weight})",
                style="List Bullet",
            )
            for run in paragraph.runs:
                run.font.size = Pt(10 * page.scale)
        recommendation = doc.add_paragraph(
            f"{_text('recommendation', locale)}: "
            f"{_recommendation_text(candidate, locale)}"
        )
        for run in recommendation.runs:
            run.font.size = Pt(10 * page.scale)

    path = Path(output)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _bounds(snapshot: SeatingSnapshot) -> tuple[int, int, int, int]:
    rows = [s.row for s in snapshot.layout.seats]
    cols = [s.col for s in snapshot.layout.seats]
    return min(rows), max(rows), min(cols), max(cols)


def _set_cell_font_size(cell, size) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = size
